from haystack.nodes import EmbeddingRetriever, BM25Retriever, PreProcessor, AnswerParser, TransformersSummarizer, PromptNode, PromptTemplate
from haystack.document_stores import PineconeDocumentStore, InMemoryDocumentStore
from haystack.utils import convert_files_to_docs, print_questions
from haystack import Pipeline
from haystack.pipelines import DocumentSearchPipeline
from typing import Annotated
from fastapi import FastAPI, Depends, HTTPException, status, Request, Form, Response, File, UploadFile
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.encoders import jsonable_encoder
from haystack import Document

from haystack.agents.memory import ConversationSummaryMemory
from haystack.agents import AgentStep, Agent, Tool
from haystack.agents.base import Agent, ToolsManager


from docx import Document as DC
from io import BytesIO
import uvicorn
import json
import os
import re
import glob
import time
import asyncio
from dotenv import load_dotenv

load_dotenv()

app = FastAPI()

origins = [
    "http://localhost.tiangolo.com",
    "https://localhost.tiangolo.com",
    "http://localhost:3000",
    "http://192.168.0.104",
    "null"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

templates = Jinja2Templates(directory="templates")

HF_TOKEN = os.getenv('HF_TOKEN')
OPEN_AI_KEY = os.getenv('OPENAI_API_KEY')


document_store = PineconeDocumentStore(api_key=os.getenv('PINECONE_API_KEY'), environment="gcp-starter",
                                similarity="dot_product",
                                embedding_dim=768)

retriever = EmbeddingRetriever(document_store = document_store,
                                embedding_model="sentence-transformers/multi-qa-mpnet-base-dot-v1")

prompt_template = PromptTemplate(prompt = """Answer the asked query based on only the given documents. If the documents do not contain the answer to the question, say that answering is not possible given the available information.
                                            Documents: {join(documents)}
                                            Query: {query}
                                            Answer: 
                                        """,
                                output_parser=AnswerParser()
                                )

prompt_node = PromptNode(
    model_name_or_path="mistralai/Mixtral-8x7B-Instruct-v0.1",
    api_key=HF_TOKEN,
    default_prompt_template=prompt_template,
    max_length=1000,
    timeout= 300,
    model_kwargs={"model_max_length": 8000}
)

query_pipeline = Pipeline()
query_pipeline.add_node(component=retriever, name="Retriever", inputs=["Query"])
query_pipeline.add_node(component=prompt_node, name="PromptNode", inputs=["Retriever"])




prompt_template_chat = PromptTemplate(prompt = """Look at the summary to get information about what has happened previously in the chat. Answer the asked query based on only the given documents and summary. Answer in less than 150 words. Your answer should only contain common spoken english words and nothing more complex than high school level english.
                                            Summary: {meta['summary']}
                                            Documents: {join(documents)}
                                            Query: {query}
                                            Answer: 
                                        """,
                                output_parser=AnswerParser()
                                )


prompt_node_chat = PromptNode(
    model_name_or_path="mistralai/Mixtral-8x7B-Instruct-v0.1",
    api_key=HF_TOKEN,
    default_prompt_template=prompt_template_chat,
    max_length=1000,
    timeout= 300,
    model_kwargs={"model_max_length": 8000}
)

chat_pipeline = Pipeline()
chat_pipeline.add_node(component=retriever, name="Retriever", inputs=["Query"])
chat_pipeline.add_node(component=prompt_node_chat, name="PromptNode", inputs=["Retriever"])


def get_additional_legal_clauses(documentType, legalClauses):

    query = f"Given a {documentType} which has the following clauses {legalClauses}. Return additional clauses that can be added to the legal agreement."

    json_response = query_pipeline.run(query=query, params={"Retriever" : {"top_k": 5}, "debug": True})

    for component in json_response['_debug']:
        print(component, json_response['_debug'][component]['exec_time_ms'])

    answers = json_response['answers']
    documents = json_response['documents']

    for ans in answers:
        answer = ans.answer
        break

    document_names = []

    for document in documents:
        document_names.append(document.id)

    return answer, document_names

def chat_with_rag_pipeline(query, chat_pipeline, summary="", legal_clauses=""):
    
    json_response = chat_pipeline.run(query=query, params={"Retriever" : {"top_k": 5}, "debug": True}, meta={'summary': legal_clauses+summary})

    for component in json_response['_debug']:
        print(component, json_response['_debug'][component]['exec_time_ms'])

    answers = json_response['answers']
    documents = json_response['documents']

    for ans in answers:
        answer = ans.answer
        break

    document_names = []

    for document in documents:
        document_names.append(document.id)

    docs = [Document(f"{summary}"), Document(f"{answer}")]

    prompt_template_summary = PromptTemplate(prompt = """Generate a historical memory such that it contains all the points in the given documents. Maintain bullet point structurings of the original documents in the summary. The summary must be such that a chatbot can identify the history of documents and answer accordingly to newer queries. The summary should be less than 100 words.
                                            Documents: {join(documents)}
                                            Answer: 
                                        """,
                                output_parser=AnswerParser()
                                )

    new_summary = prompt_node.prompt(prompt_template=prompt_template_summary, documents=docs)[0].answer

    return answer, document_names, new_summary

########################################################################################################################
## Creating a common document store for contract templates uploaded at initial load

document_store_for_template = None

async def get_document_store(file=None):

    global document_store_for_template
    # print(document_store_for_template)
    if document_store_for_template is not None:
        return document_store_for_template
    else:
        filedir = 'docs'
        with open(f"{filedir}/{file.filename}", "wb") as f:
            content = await file.read()
            f.write(content)

        docs = convert_files_to_docs(dir_path=filedir)

        document_store_for_template = InMemoryDocumentStore(use_bm25=True)

        preprocessor = PreProcessor(
            clean_empty_lines=True,
            clean_whitespace=True,
            clean_header_footer=True,
            split_by="sentence",
            split_length=5,
            split_overlap=2,
            split_respect_sentence_boundary=False,
        )

        preprocessed_docs = preprocessor.process(docs)

        document_store_for_template.write_documents(preprocessed_docs)

        files = glob.glob(f"{filedir}/*")
        for f in files:
            os.remove(f)

        return document_store_for_template

def replace_dots_with_mask(input_string):

    counter = 1
    result_string = ''
    start = 0
    for m in re.finditer(r'\.{3,}|_{3,}', input_string):
        end, newstart = m.span()
        result_string += input_string[start:end]
        rep = f" [MASK{counter}] "
        result_string += rep
        start = newstart
        counter += 1
    result_string += input_string[start:]

    return result_string.strip()

def clean_json_string(input_string):
    # Filter out invalid characters
    valid_chars = [char for char in input_string if char.isprintable() and ord(char) < 128]

    # Reconstruct the cleaned string
    cleaned_string = ''.join(valid_chars)

    cleaned_string = re.sub(r'[^\x20-\x7E]', '', cleaned_string)

    return cleaned_string

########################################################################################################################
## Route definitions of the API

@app.get("/")
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/relevant-documents")
async def index(relevant_documents: str = Form(...)):

    relevant_documents_ids = json.loads(relevant_documents)
    relevant_documents = document_store_for_template.get_documents_by_id(ids=relevant_documents_ids)

    content = []

    for relevant_doc in relevant_documents:
        content.append(relevant_doc.content)

    response_data = jsonable_encoder(json.dumps({"relevant_documents_content": content}))

    res = Response(response_data)
    return res




@app.post("/legal-clause-gen")
async def get_clauses(document_type: Annotated[str, Form()], legal_clauses: Annotated[str, Form()]):
    start_time = time.time()
    answer, relevant_documents = get_additional_legal_clauses(document_type, legal_clauses)
    response_data = jsonable_encoder(json.dumps({"answer": answer, "relevant_documents": relevant_documents}))
    print("Time took to process the request and return response is {} sec".format(time.time() - start_time))
    res = Response(response_data)
    return res

@app.post("/legal-clause-chat")
async def chat_clauses(summary: Annotated[str, Form()], query: Annotated[str, Form()], legal_clauses: Annotated[str, Form()]):
    start_time = time.time()
    answer, relevant_documents, updated_summary = chat_with_rag_pipeline(query, chat_pipeline, summary, legal_clauses)
    response_data = jsonable_encoder(json.dumps({"answer": answer, "relevant_documents": relevant_documents, "summary": updated_summary}))
    print("Time took to process the request and return response is {} sec".format(time.time() - start_time))
    res = Response(response_data)
    return res

@app.post("/doc-chat")
async def chat_doc(summary: Annotated[str, Form()], query: Annotated[str, Form()]):
    start_time = time.time()

    document_store = await get_document_store()

    retriever = BM25Retriever(document_store=document_store, top_k=10)

    chat_pipeline_solo = Pipeline()
    chat_pipeline_solo.add_node(component=retriever, name="Retriever", inputs=["Query"])
    chat_pipeline_solo.add_node(component=prompt_node_chat, name="PromptNode", inputs=["Retriever"])

    answer, relevant_documents, updated_summary = chat_with_rag_pipeline(query, chat_pipeline_solo, summary)
    response_data = jsonable_encoder(json.dumps({"answer": answer, "relevant_documents": relevant_documents, "summary": updated_summary}))
    print("Time took to process the request and return response is {} sec".format(time.time() - start_time))
    res = Response(response_data)
    return res

    
@app.post("/doc-upload")
async def analyze_doc(file: UploadFile = File(...)):
    start_time = time.time()
    
    global document_store_for_template

    document_store_for_template = None

    document_store = await get_document_store(file)

    retriever = BM25Retriever(document_store=document_store, top_k=10)

    prompt_template = PromptTemplate(
        prompt="""Synthesize a comprehensive answer from the following text for the given question.
                                Provide a clear and concise response that summarizes the key points and information presented in the text.
                                \n\n Related text: {join(documents)} \n\n Question: {query} \n\n Answer:""",
        output_parser=AnswerParser(),
    )

    prompt_node = PromptNode(        
        model_name_or_path="mistralai/Mixtral-8x7B-Instruct-v0.1",
        api_key=HF_TOKEN,
        default_prompt_template=prompt_template,
        max_length=1000,
        timeout= 300,
        model_kwargs={"model_max_length": 5000}
    )

    pipe = Pipeline()
    pipe.add_node(component=retriever, name="retriever", inputs=["Query"])
    pipe.add_node(component=prompt_node, name="prompt_node", inputs=["retriever"])

    documentType = pipe.run(query="Identify the type of legal contract in the given context. Answer only the type of legal contract and enclose it in curly braces. Add no other extra words to the answer.")

    # print(documentType["answers"][0].answer)

    legalClauses = pipe.run(query="What legal clauses are present in the given legal document.")

    # print(legalClauses["answers"][0].answer)

    response_data = jsonable_encoder(json.dumps({"documentType": documentType["answers"][0].answer, "legalClauses": legalClauses["answers"][0].answer}))
    res = Response(response_data)

    print("Time took to process the request and return response is {} sec".format(time.time() - start_time))

    return res


@app.post("/question-gen")
async def question_gen(file: UploadFile = File(...)):

    filedir = 'docs'
    with open(f"{filedir}/{file.filename}", "wb") as f:
        content = await file.read()
        f.write(content)

    docs = convert_files_to_docs(dir_path=filedir)

    preprocessor = PreProcessor(
        clean_empty_lines=True,
        clean_whitespace=True,
        clean_header_footer=True,
        split_by="sentence",
        split_length=1,
        split_respect_sentence_boundary=False,
    )

    preprocessed_docs = preprocessor.process(docs)

    files = glob.glob(f"{filedir}/*")
    for f in files:
        os.remove(f)

    prompt_template = PromptTemplate(
    prompt="""
    Generate a single, clear question for each [MASK1], [MASK2], [MASK3], etc., in the document. Ensure that each question is uniquely tailored to gather information for the specific [MASK] it corresponds to. Do not generate any extra questions unrelated to the specified placeholders. Return the answer as a json object where the keys are the [MASK] tokens and the values are the questions. 
    Document: {document}
    """,
    output_parser=AnswerParser(),
    )


    prompt_node = PromptNode(        
        model_name_or_path="gpt-3.5-turbo",
        api_key=OPEN_AI_KEY,
        # model_name_or_path="mistralai/Mixtral-8x7B-Instruct-v0.1",
        # api_key=HF_TOKEN,
        default_prompt_template=prompt_template,
        # output_variable="my_answer",
        model_kwargs={"temperature": 0.2, "top_p": 0.1, "max_tokens": 2000}
    )

    quesDocsIndex = []

    for idx, document in enumerate(preprocessed_docs):
        blankspaces = re.findall(r'\.{3,}|_{3,}', document.content)

        if not blankspaces:
            continue

        quesDocsIndex.append(idx)

    totalDocs = len(quesDocsIndex)

    async def generate_ques():
        for idx, quesDocIndex in enumerate(quesDocsIndex):

            document = preprocessed_docs[quesDocIndex].content

            document = replace_dots_with_mask(document)

            document = document.replace("\u2026", "")

            result = prompt_node.prompt(prompt_template=prompt_template, document=document)

            result[0].answer = re.sub(r',(\s*})', r'\1', result[0].answer)
            answerJson = json.loads(str(result[0].answer))
            removeItems = []
            for key, value in answerJson.items():
                if key not in document:
                    removeItems.append(key)
            
            for key in removeItems:
                answerJson.pop(key)

            print('{"questions": ' + str(json.dumps(answerJson)) + ',"document": "' + document + '"}')

            document = clean_json_string(document)

            yield '{"questions": ' + str(json.dumps(answerJson)) + ',"document": "' + str(document.replace('\n', '\\n')) + '","docIndex":'+ str(idx) +',"totalDocs": ' + str(totalDocs) + '}'



        # for i in range(5):
        #     yield f"data: {{'key': '{preprocessed_docs[0].content}'}}\n\n"
        #     await asyncio.sleep(5)

    return StreamingResponse(generate_ques(), media_type='text/event-stream')


@app.post("/download-docx")
async def download_docx(substitution_strings: str = Form(...), file: UploadFile = File(...)):
    # Read the uploaded docx file
    content = await file.read()
    
    substitution_strings = json.loads(substitution_strings)

    try:
        # Create a docx document from the content
        doc = DC(BytesIO(content))
    except Exception as e:
        return {"error": f"Error decoding docx file: {str(e)}"}
    
    # Define your regex pattern and multiple substitution strings
    regex_pattern = r'\.{3,}|_{3,}'

    substitution_counter = 0
    
    
    # Process each paragraph in the document
    for paragraph in doc.paragraphs:
        input_string = paragraph.text
        result_string = ''
        start = 0
        for m in re.finditer(regex_pattern, input_string):
            end, newstart = m.span()
            result_string += input_string[start:end]
            if substitution_counter >= len(substitution_strings):
                rep = '[MISSED]'
            else:
                rep = f" {substitution_strings[substitution_counter]} "
            result_string += rep
            start = newstart
            substitution_counter += 1
        result_string += input_string[start:]
        paragraph.text = result_string
        paragraph.text = paragraph.text.replace("\u2026", "")

        
    # Process each table in the document
    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             cell.text = re.sub(regex_pattern, get_next_substitution(), cell.text)
    
    # Save the modified document to a BytesIO object
    output_doc = BytesIO()
    doc.save(output_doc)
    output_doc.seek(0)
    
    # Return the modified document as a response
    return StreamingResponse(output_doc, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=output.docx"})
if __name__ == "__main__":
    uvicorn.run("app:app", host='0.0.0.0', port=8001, reload=True)



